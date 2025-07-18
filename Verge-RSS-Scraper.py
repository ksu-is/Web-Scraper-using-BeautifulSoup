# With extensive help from ChatGPT (since I don't know all these libraries) I modified the old VergeScraper program to get links from RSS feeds.  'We' had to work on the hyperlink function as it made the URLs look like plain unclickable text which would make the reader think there was no link to click in the Word doc.  I worked it out so that the URLs use standard link formatting, i.e. blue, underlined text.

import feedparser
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from datetime import datetime

# Function to add a hyperlink to a paragraph
# I had to troubleshoot the hyperlink function as it made the URLs look like plain, unclickable text which would make the reader think there was no link to click in the Word doc.  I worked it out so that the URLs use standard link formatting, i.e. blue, underlined text.
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, reltype='http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set style manually: blue and underlined
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '3cffd0') # match some of the color scheme from TheVerge.com
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# Parse The Verge RSS feed
feed_url = 'https://www.theverge.com/rss/index.xml'
feed = feedparser.parse(feed_url)
entries = feed.entries[:10]

# Prepare output Word document
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
filename = f"theverge_headlines_{timestamp}.docx"
doc = Document()
doc.add_heading('Today on The Verge (RSS)', level=1) # better heading
doc.add_paragraph(f"Scraped on {timestamp}\n")

# Populate document with URLs shown and hyperlinked
if not entries:
    doc.add_paragraph("No articles found in the RSS feed.")
else:
    for entry in entries:
        p = doc.add_paragraph("\u2022 ")
        add_hyperlink(p, entry.link, entry.link)
        print(entry.title, "-", entry.link)

# Save the Word document
doc.save(filename)
print(f"Saved to {filename}")
