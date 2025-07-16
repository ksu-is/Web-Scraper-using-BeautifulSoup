import requests
from bs4 import BeautifulSoup
# import csv
from docx import Document # testing with exporting to Word instead of csv
from datetime import datetime
import re

# Fetch the page content
url = 'https://www.theverge.com'
response = requests.get(url)
html_source = response.text
soup = BeautifulSoup(html_source, 'html.parser')

# Find and filter links that match article patterns
all_links = soup.find_all('a', href=True)
article_links = []
seen = set()

for link in all_links:
    href = link['href']
    text = link.get_text(strip=True)

    if not text or href in seen:
        continue
    seen.add(href)

    # Match links like "/section/123456/title"
    if re.match(r"^/[\w\-]+/\d{6,}/", href):
        full_url = f"https://www.theverge.com{href}" if not href.startswith("http") else href
        article_links.append((text, full_url))

# Create the CSV/Word filename with timestamp
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
# filename = f"theverge_headlines_{timestamp}.csv"
filename = f"theverge_headlines_{timestamp}.docx"

# Write to CSV
# with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
#     writer = csv.writer(csvfile)
#     writer.writerow(['Headline', 'URL'])

#     for headline, url in article_links[:10]:  # Only first 10
#         writer.writerow([headline, url])
#         print(f"{headline} - {url}")

# Write to Word
doc = Document()
doc.add_heading('The Verge Headlines', level=1)
doc.add_paragraph(f"Scraped on {timestamp}\n")

for headline, url in article_links[:10]:  # Only first 10
    doc.add_paragraph(f"{headline}\n{url}\n")
    print(f"{headline} - {url}")

# Next, I'll try to add hyperlinks within the Word doc, or output only the article title with a hyperlink
doc.save(filename)

if not article_links:
    print("No article links found. The page structure may have changed.")
